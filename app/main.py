import html
import io
import json
import logging
import os
import re
import shutil
import time
import uuid
from contextlib import asynccontextmanager
from pathlib import Path

from docx import Document
from docx.shared import Pt
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from openai import OpenAI
from pydantic import BaseModel, Field

from .parsers import detect_file_type, parse_document

logging.basicConfig(
    format="%(asctime)s %(levelname)s %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

ROOT = Path(__file__).resolve().parent.parent
UPLOADS_DIR = ROOT / "uploads"
LOGS_DIR = ROOT / "logs"
PROMPTS_FILE = ROOT / "prompts.json"
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)


def _load_prompts() -> dict:
    """Load prompts configuration from JSON file."""
    default_prompts = {
        "system": [
            "You are a precise translation engine.",
            "Translate the given text into {target_language}.",
            "Output ONLY the translation text.",
            "Do not include explanations, quotes, commentary, prefixes, or extra lines.",
            "Maintain consistent terminology and style with the prior translated passages provided.",
        ],
        "context_header": "=== PREVIOUSLY TRANSLATED PASSAGES (for context and consistency) ===",
        "context_entry": '{index}. "{original}" → "{translation}"',
        "user_prompt": "=== NEW TEXT TO TRANSLATE ===\n{text}",
        "temperature": 0.1,
    }
    
    if PROMPTS_FILE.exists():
        try:
            with open(PROMPTS_FILE, "r", encoding="utf-8") as f:
                loaded = json.load(f)
                # Merge with defaults (so missing keys use defaults)
                return {**default_prompts, **loaded}
        except Exception as e:
            logger.warning(f"Failed to load prompts.json: {e}, using defaults")
    
    return default_prompts


PROMPTS = _load_prompts()


# ============================================================================
# Environment Settings
# ============================================================================

def _get_env(name: str, default: str = "") -> str:
    """Get environment variable, returning default if empty or not set."""
    val = os.getenv(name, "")
    return val.strip() if val.strip() else default


def _get_env_int(name: str, default: int) -> int:
    """Get environment variable as integer."""
    val = _get_env(name, str(default))
    try:
        return int(val)
    except ValueError:
        return default


def _mask_api_key(key: str) -> str:
    """Mask API key for display, showing only first 4 and last 4 chars."""
    if not key or len(key) < 12:
        return "••••••••" if key else ""
    return f"{key[:4]}••••••••{key[-4:]}"


ENV_API_KEY = _get_env("API_KEY")
ENV_API_ENDPOINT = _get_env("API_ENDPOINT", "https://api.openai.com/v1")
ENV_MODEL = _get_env("MODEL", "gpt-4o-mini")
MAX_UPLOAD_MB = _get_env_int("MAX_UPLOAD_MB", 50)
MAX_FILE_SIZE = MAX_UPLOAD_MB * 1024 * 1024
CLEANUP_LOGS_DAYS = _get_env_int("CLEANUP_LOGS_DAYS", 30)


# ============================================================================
# Cleanup Functions
# ============================================================================

def _clear_uploads() -> None:
    """Remove all files from uploads directory."""
    if not UPLOADS_DIR.exists():
        return
    count = sum(1 for f in UPLOADS_DIR.iterdir() if f.is_file())
    if count > 0:
        shutil.rmtree(UPLOADS_DIR)
        UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
        logger.info(f"Cleared {count} uploaded file(s)")


def _cleanup_old_logs() -> None:
    """Delete log files older than CLEANUP_LOGS_DAYS."""
    if CLEANUP_LOGS_DAYS <= 0 or not LOGS_DIR.exists():
        return
    
    cutoff = time.time() - (CLEANUP_LOGS_DAYS * 24 * 60 * 60)
    deleted = 0
    
    for file_path in LOGS_DIR.iterdir():
        if file_path.is_file():
            try:
                if file_path.stat().st_mtime < cutoff:
                    file_path.unlink()
                    deleted += 1
            except OSError:
                pass
    
    if deleted > 0:
        logger.info(f"Cleaned {deleted} log file(s) older than {CLEANUP_LOGS_DAYS} days")


@asynccontextmanager
async def lifespan(app):
    """Handle startup and shutdown cleanup."""
    logger.info("Starting up...")
    _clear_uploads()
    _cleanup_old_logs()
    yield
    logger.info("Shutting down...")
    _clear_uploads()


def _normalize_whitespace(s: str, preserve_paragraphs: bool = False) -> str:
    """Collapse whitespace into single spaces.
    
    If preserve_paragraphs is True, double newlines (paragraph breaks) are preserved.
    """
    if preserve_paragraphs:
        # Split on double newlines (paragraph breaks), normalize each paragraph, rejoin
        paragraphs = re.split(r'\n\s*\n', s)
        normalized = [re.sub(r'\s+', ' ', p).strip() for p in paragraphs]
        return '\n\n'.join(p for p in normalized if p)
    return re.sub(r"\s+", " ", s).strip()


def _target_language_label(lang: str) -> str:
    """Normalize and validate target language."""
    lang_norm = (lang or "").strip()
    if not lang_norm:
        return "English"
    abbrev = {
        "en": "English", "zh": "Chinese (Simplified)", "es": "Spanish", "fr": "French",
        "de": "German", "ja": "Japanese", "ko": "Korean", "ru": "Russian",
        "pt": "Portuguese", "it": "Italian", "zh-cn": "Chinese (Simplified)",
        "zh-hans": "Chinese (Simplified)", "zh-hant": "Chinese (Traditional)",
    }
    lower = lang_norm.lower()
    if lower in abbrev:
        return abbrev[lower]
    return lang_norm


app = FastAPI(title="Step Translate", lifespan=lifespan)


# ============================================================================
# Models
# ============================================================================

class PassageModel(BaseModel):
    id: str
    text: str
    page: int | None = None
    style: str = "paragraph"  # 'title', 'heading', 'author', 'paragraph'


class UploadResponse(BaseModel):
    file_id: str
    filename: str
    file_type: str
    passages: list[PassageModel]
    pdf_url: str | None = None  # For PDFs, provide URL to render with PDF.js


class PriorTranslation(BaseModel):
    translation: str


class TranslateRequest(BaseModel):
    selected_text: str = Field(min_length=1)
    target_language: str = Field(description="Target language for translation")
    prior_translations: list[PriorTranslation] = Field(default_factory=list)
    api_key: str | None = None
    api_endpoint: str | None = None
    temperature: float | None = None
    system_prompt: str | None = None
    user_prompt: str | None = None
    model: str | None = None


class TranslateResponse(BaseModel):
    translation: str


class AdaptRequest(BaseModel):
    selected_text: str = Field(min_length=1)
    target_language: str = Field(description="Target language for adaptation")
    additional_instructions: str | None = None
    api_key: str | None = None
    api_endpoint: str | None = None
    temperature: float | None = None
    model: str | None = None
    adapt_system_prompt: str | None = None
    adapt_user_prompt: str | None = None


class AdaptResponse(BaseModel):
    adapted_text: str


class ExportRequest(BaseModel):
    text: str = Field(min_length=1)
    format: str = Field(pattern="^(txt|docx|pdf)$")
    filename: str = "translation"


# ============================================================================
# Endpoints
# ============================================================================

@app.post("/api/upload", response_model=UploadResponse)
async def upload_document(file: UploadFile = File(...)) -> UploadResponse:
    """Upload a document (PDF, TXT, or DOCX) and extract passages."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="Missing filename.")
    
    # Detect file type
    file_type = detect_file_type(file.filename, file.content_type)
    if not file_type:
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Please upload PDF, TXT, or DOCX."
        )
    
    # Read file data
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file.")
    if len(data) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large (max {MAX_UPLOAD_MB}MB).")
    
    # Validate PDF header if PDF
    if file_type == 'pdf' and not data.startswith(b"%PDF"):
        raise HTTPException(status_code=400, detail="Invalid PDF file.")
    
    # Save file
    file_id = str(uuid.uuid4())
    ext = Path(file.filename).suffix.lower() or f".{file_type}"
    out_path = UPLOADS_DIR / f"{file_id}{ext}"
    out_path.write_bytes(data)
    
    # Parse document into passages
    try:
        passages = parse_document(out_path, file_type)
    except Exception as e:
        logger.error(f"Failed to parse document: {e}")
        out_path.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail=f"Failed to parse document: {str(e)}")
    
    if not passages:
        out_path.unlink(missing_ok=True)
        raise HTTPException(status_code=400, detail="No text content found in document.")
    
    logger.info(f"Uploaded {file.filename} ({file_type}): {len(passages)} passages")
    
    # For PDFs, provide URL for PDF.js rendering
    pdf_url = f"/uploads/{file_id}{ext}" if file_type == "pdf" else None
    
    return UploadResponse(
        file_id=file_id,
        filename=file.filename,
        file_type=file_type,
        passages=[
            PassageModel(id=p.id, text=p.text, page=p.page, style=p.style)
            for p in passages
        ],
        pdf_url=pdf_url,
    )


@app.post("/api/translate", response_model=TranslateResponse)
async def translate(req: TranslateRequest) -> TranslateResponse:
    """Translate text with context from prior translations."""
    # Get API settings: use request values if provided, otherwise fall back to env vars
    api_key = (req.api_key or "").strip() or ENV_API_KEY
    if not api_key:
        raise HTTPException(
            status_code=400,
            detail="API key is required. Configure it in Settings or set API_KEY environment variable."
        )
    
    api_endpoint = (req.api_endpoint or "").strip() or ENV_API_ENDPOINT
    model = (req.model or "").strip() or ENV_MODEL
    target = _target_language_label(req.target_language)
    
    # Normalize the selected text (preserve paragraph breaks for multi-passage translations)
    selected = _normalize_whitespace(req.selected_text, preserve_paragraphs=True)
    if not selected:
        raise HTTPException(status_code=400, detail="No text to translate.")
    
    # Build system prompt: use request override or fall back to config file
    if req.system_prompt:
        system = req.system_prompt.format(target_language=target)
    else:
        system_lines = [
            line.format(target_language=target) for line in PROMPTS["system"]
        ]
        system = "\n".join(system_lines)
    
    # Build user message with context
    user_parts = []
    
    if req.prior_translations:
        user_parts.append(PROMPTS["context_header"])
        for i, pt in enumerate(req.prior_translations[-5:], 1):  # Last 5 translations max
            trans = _normalize_whitespace(pt.translation)[:400]
            user_parts.append(f"[{i}] {trans}")
        user_parts.append("")
    
    # Use request override for user prompt or fall back to config
    user_prompt_template = req.user_prompt if req.user_prompt else PROMPTS["user_prompt"]
    user_parts.append(user_prompt_template.format(text=selected))
    user = "\n".join(user_parts)
    
    # Temperature: request override > config file > default
    temperature = req.temperature if req.temperature is not None else PROMPTS.get("temperature", 0.1)
    
    try:
        client = OpenAI(api_key=api_key, base_url=api_endpoint, timeout=60.0)
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=temperature,
        )
    except Exception as e:
        logger.error(f"Translation API error: {e}")
        raise HTTPException(status_code=502, detail=f"Translation API error: {str(e)}")
    
    text = (resp.choices[0].message.content or "").strip()
    if not text:
        raise HTTPException(status_code=502, detail="Empty translation response.")
    
    logger.info(f"TRANSLATE [{target}] Input: {selected[:100]}...")
    logger.info(f"TRANSLATE [{target}] Output: {text[:100]}...")
    
    return TranslateResponse(translation=text)


@app.post("/api/adapt", response_model=AdaptResponse)
async def adapt_text(req: AdaptRequest) -> AdaptResponse:
    """Adapt/naturalize translated text to sound more native in the target language."""
    api_key = (req.api_key or "").strip() or ENV_API_KEY
    if not api_key:
        raise HTTPException(
            status_code=400,
            detail="API key is required. Configure it in Settings or set API_KEY environment variable."
        )

    api_endpoint = (req.api_endpoint or "").strip() or ENV_API_ENDPOINT
    model = (req.model or "").strip() or ENV_MODEL
    target = _target_language_label(req.target_language)

    selected = _normalize_whitespace(req.selected_text, preserve_paragraphs=True)
    if not selected:
        raise HTTPException(status_code=400, detail="No text to adapt.")

    # Build system prompt
    if req.adapt_system_prompt:
        system = req.adapt_system_prompt.format(target_language=target)
    else:
        adapt_system_lines = PROMPTS.get("adapt_system", [
            "You are a skilled editor for {target_language} text.",
            "Rewrite the following translated passage so it reads naturally in {target_language}.",
            "Preserve the original meaning but improve fluency and idiom.",
            "Output ONLY the adapted text.",
        ])
        system = "\n".join(line.format(target_language=target) for line in adapt_system_lines)

    # Build user prompt
    adapt_user_template = req.adapt_user_prompt if req.adapt_user_prompt else PROMPTS.get("adapt_user_prompt", "=== TEXT TO ADAPT ===\n{text}")
    user = adapt_user_template.format(text=selected)

    # Append additional instructions if provided
    extra = (req.additional_instructions or "").strip()
    if extra:
        user += f"\n\n=== ADDITIONAL INSTRUCTIONS ===\n{extra}"

    temperature = req.temperature if req.temperature is not None else PROMPTS.get("temperature", 0.1)

    try:
        client = OpenAI(api_key=api_key, base_url=api_endpoint, timeout=60.0)
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=temperature,
        )
    except Exception as e:
        logger.error(f"Adapt API error: {e}")
        raise HTTPException(status_code=502, detail=f"Adapt API error: {str(e)}")

    text = (resp.choices[0].message.content or "").strip()
    if not text:
        raise HTTPException(status_code=502, detail="Empty adaptation response.")

    logger.info(f"ADAPT [{target}] Input: {selected[:100]}...")
    logger.info(f"ADAPT [{target}] Output: {text[:100]}...")

    return AdaptResponse(adapted_text=text)


@app.post("/api/export")
async def export_translation(req: ExportRequest) -> StreamingResponse:
    """Export translation text as DOCX or PDF."""
    text = req.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="No text to export.")
    
    safe_filename = re.sub(r'[^\w\-.]', '_', req.filename)[:100]
    
    if req.format == "docx":
        # Create DOCX document
        doc = Document()
        
        # Split by double newlines to create paragraphs
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                para = doc.add_paragraph()
                run = para.add_run(para_text)
                run.font.size = Pt(11)
        
        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_filename}.docx"'},
        )
    
    elif req.format == "pdf":
        # Create PDF using reportlab
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        except ImportError:
            raise HTTPException(
                status_code=500, 
                detail="PDF export requires reportlab. Install with: pip install reportlab"
            )
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
        )
        
        # Use bundled Noto Sans SC TTF font for CJK support
        font_name = "Helvetica"
        try:
            font_path = Path(__file__).parent / "fonts" / "NotoSansSC-Regular.ttf"
            
            if not font_path.exists():
                raise FileNotFoundError(f"Font file not found at: {font_path}")
            
            pdfmetrics.registerFont(TTFont("NotoSansSC", str(font_path)))
            font_name = "NotoSansSC"
            logger.info(f"Loaded CJK font from: {font_path}")
        except Exception as e:
            logger.warning(f"Could not load CJK fonts: {e}, falling back to Helvetica")
        
        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            'Body',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=11,
            leading=16,
            spaceAfter=12,
        )
        
        story = []
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # Escape XML special chars and preserve line breaks
                para_text = para_text.replace('&', '&amp;')
                para_text = para_text.replace('<', '&lt;')
                para_text = para_text.replace('>', '&gt;')
                para_text = para_text.replace('\n', '<br/>')
                story.append(Paragraph(para_text, body_style))
                story.append(Spacer(1, 6))
        
        doc.build(story)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{safe_filename}.pdf"'},
        )
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use txt, docx, or pdf.")


# ============================================================================
# Static file serving
# ============================================================================

static_dir = ROOT / "static"
_index_html_path = static_dir / "index.html"
_cache_version = uuid.uuid4().hex[:8]

# Pre-compute server defaults for injection into HTML
def _build_server_defaults_script() -> str:
    """Build a script tag with server defaults for the frontend."""
    defaults = {
        "apiEndpoint": ENV_API_ENDPOINT,
        "model": ENV_MODEL,
        "temperature": PROMPTS.get("temperature", 0.1),
        "systemPrompt": "\n".join(PROMPTS["system"]),
        "userPrompt": PROMPTS["user_prompt"],
        "adaptSystemPrompt": "\n".join(PROMPTS.get("adapt_system", [])),
        "adaptUserPrompt": PROMPTS.get("adapt_user_prompt", "=== TEXT TO ADAPT ===\n{text}"),
        "hasApiKey": bool(ENV_API_KEY),
        "apiKeyMasked": _mask_api_key(ENV_API_KEY),
    }
    # Escape for safe embedding in HTML
    json_str = json.dumps(defaults, ensure_ascii=False)
    escaped = html.escape(json_str, quote=False)
    return f'<script>window.SERVER_DEFAULTS = {escaped};</script>'

_server_defaults_script = _build_server_defaults_script()


@app.get("/", response_class=HTMLResponse)
async def serve_index() -> HTMLResponse:
    """Serve index.html with server defaults and cache-busting version."""
    html_content = _index_html_path.read_text()
    html_content = html_content.replace("{{CACHE_VERSION}}", _cache_version)
    html_content = html_content.replace("{{SERVER_DEFAULTS}}", _server_defaults_script)
    return HTMLResponse(content=html_content)


# Serve uploaded files (for PDF.js to access PDFs)
app.mount("/uploads", StaticFiles(directory=UPLOADS_DIR), name="uploads")

app.mount("/", StaticFiles(directory=static_dir, html=True), name="static")
